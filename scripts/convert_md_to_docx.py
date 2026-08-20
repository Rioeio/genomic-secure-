import os
import re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def latex_to_pretty_math(latex_str: str) -> str:
    """
    Transforms LaTeX equations into beautifully formatted mathematical notation.
    """
    s = latex_str.strip()
    if s.startswith('$$') and s.endswith('$$'):
        s = s[2:-2].strip()
    elif s.startswith('$') and s.endswith('$'):
        s = s[1:-1].strip()

    # Step 1: Text blocks
    s = re.sub(r'\\text\{([^}]+)\}', r'\1', s)
    
    # Step 2: Fractions (handle nested)
    def clean_frac(match):
        num = match.group(1).strip()
        den = match.group(2).strip()
        if len(num) > 1 and not (num.startswith('(') and num.endswith(')')):
            num_str = f"{num}"
        else:
            num_str = num
        if len(den) > 1 and not (den.startswith('(') and den.endswith(')')):
            den_str = f"{den}"
        else:
            den_str = den
        return f"({num_str} / {den_str})"

    for _ in range(5):
        s = re.sub(r'\\frac\{([^{}]+)\}\{([^{}]+)\}', clean_frac, s)

    # Step 3: Greek & Special Symbols
    symbols = [
        (r'\\mathcal\{L\}', 'ℒ'),
        (r'\\mathcal\{M\}', 'ℳ'),
        (r'\\mathcal\{D\}', '𝒟'),
        (r'\\mathcal\{Q\}', '𝒬'),
        (r'\\mathcal\{R\}', 'ℛ'),
        (r'\\mathbb\{R\}', 'ℝ'),
        (r'\\mathbb\{I\}', '𝕀'),
        (r'\\mathbf\{1\}', '𝟏'),
        (r'\\nabla_w', '∇_w'),
        (r'\\nabla_b', '∇_b'),
        (r'\\nabla', '∇'),
        (r'\\epsilon', 'ε'),
        (r'\\delta', 'δ'),
        (r'\\alpha', 'α'),
        (r'\\lambda', 'λ'),
        (r'\\sigma', 'σ'),
        (r'\\beta', 'β'),
        (r'\\Phi', 'Φ'),
        (r'\\Delta_S', 'Δ_S'),
        (r'\\Delta', 'Δ'),
        (r'\\widetilde\{\\Delta w\}', 'Δ̃w'),
        (r'\\widetilde\{\\Delta W\}', 'Δ̃W'),
        (r'\\sum_\{i=1\}\^\{n_k\}', '∑(i=1…nₖ)'),
        (r'\\sum_\{j=1\}\^K', '∑(j=1…K)'),
        (r'\\sum_\{k=1\}\^K', '∑(k=1…K)'),
        (r'\\sum_\{t=1\}\^T', '∑(t=1…T)'),
        (r'\\sum', '∑'),
        (r'\\in', ' ∈ '),
        (r'\\ge', ' ≥ '),
        (r'\\le', ' ≤ '),
        (r'\\approx', ' ≈ '),
        (r'\\times', ' × '),
        (r'\\cdot', ' · '),
        (r'\\odot', ' ⊙ '),
        (r'\\to', ' → '),
        (r'\\top', 'ᵀ'),
        (r'\\|', '‖'),
        (r'\\ln', 'ln'),
        (r'\\min', 'min'),
        (r'\\max', 'max'),
        (r'\\left\(', '('),
        (r'\\right\)', ')'),
        (r'\\left\[', '['),
        (r'\\right\]', ']'),
        (r'\\left\\\{', '{'),
        (r'\\right\\\}', '}'),
        (r'\\\{', '{'),
        (r'\\\}', '}'),
    ]

    for pat, rep in symbols:
        s = re.sub(pat, rep, s)

    # Subscripts and superscripts cleanup
    sub_sup = [
        (r'\^\{2\}', '²'),
        (r'\^2', '²'),
        (r'_2\^2', '₂²'),
        (r'\^\{\(t\+1\)\}', 'ᵗ⁺¹'),
        (r'\^\{\(t\)\}', 'ᵗ'),
        (r'\^\{\(T\)\}', 'ᵀ'),
        (r'\^\{\(k\)\}', 'ᵏ'),
        (r'_k', 'ₖ'),
        (r'_i', 'ᵢ'),
        (r'_j', 'ⱼ'),
        (r'_\{pos\}', '_pos'),
        (r'_\{weight\}', '_weight'),
        (r'_\{global\}', '_global'),
        (r'_\{local\}', '_local'),
        (r'_\{trained\}', '_trained'),
        (r'_\{pop\}', '_pop'),
        (r'_\{step\}', '_step'),
        (r'_\{total\}', '_total'),
    ]
    for pat, rep in sub_sup:
        s = re.sub(pat, rep, s)

    # Final cleanup of stray backslashes and brackets
    s = s.replace('\\', '')
    s = re.sub(r'\{([a-zA-Z0-9_]+)\}', r'\1', s)
    s = re.sub(r'\s+', ' ', s).strip()
    return s

def add_math_callout_box(doc, latex_eq: str):
    """
    Creates a dedicated, beautifully formatted equation card in Microsoft Word.
    """
    pretty_eq = latex_to_pretty_math(latex_eq)
    
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    
    # Modern subtle mathematical background (ice blue / slate tint)
    set_cell_background(cell, "F0F4F8")
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    
    # Equation text in Cambria Math font
    run = p.add_run(pretty_eq)
    run.font.name = 'Cambria Math'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D) # Deep professional navy
    
    # Add small spacing after equation
    doc.add_paragraph()

def render_paragraph_with_inline_math(paragraph, text: str):
    """
    Parses mixed text containing bold (**text**), inline code (`code`), and inline math ($math$).
    """
    # Tokenize by bold, code, and inline math
    pattern = r'(\*\*.*?\*\*|`.*?`|\$.*?\$)'
    tokens = re.split(pattern, text)
    
    for token in tokens:
        if not token:
            continue
        if token.startswith('**') and token.endswith('**') and len(token) >= 4:
            r = paragraph.add_run(token[2:-2])
            r.font.bold = True
        elif token.startswith('`') and token.endswith('`') and len(token) >= 2:
            r = paragraph.add_run(token[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0x80, 0x5A, 0xD5) # Purple for code
        elif token.startswith('$') and token.endswith('$') and len(token) >= 2:
            math_text = latex_to_pretty_math(token)
            r = paragraph.add_run(math_text)
            r.font.name = 'Cambria Math'
            r.font.size = Pt(10.5)
            r.font.italic = True
            r.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D) # Navy for math
        else:
            paragraph.add_run(token)

def markdown_to_docx_with_equations(md_path: str, docx_path: str, doc_title: str):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = docx.Document()
    
    # Page setup (1-inch margins)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base styling
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)

    i = 0
    in_code_block = False
    code_lines = []
    
    while i < len(lines):
        line = lines[i].rstrip('\r\n')

        # 1. Standalone Equation Blocks ($$...$$)
        if line.strip().startswith('$$'):
            eq_lines = [line.strip()]
            if not (line.strip().endswith('$$') and len(line.strip()) > 2):
                i += 1
                while i < len(lines) and not lines[i].strip().endswith('$$'):
                    eq_lines.append(lines[i].strip())
                    i += 1
                if i < len(lines):
                    eq_lines.append(lines[i].strip())
            full_eq = ' '.join(eq_lines)
            add_math_callout_box(doc, full_eq)
            i += 1
            continue
            
        # 2. Code blocks (```...```)
        if line.startswith('```'):
            if in_code_block:
                code_text = '\n'.join(code_lines)
                table = doc.add_table(rows=1, cols=1)
                table.autofit = False
                table.columns[0].width = Inches(6.5)
                cell = table.cell(0, 0)
                set_cell_background(cell, "F7FAFC")
                cp = cell.paragraphs[0]
                cp.paragraph_format.space_before = Pt(4)
                cp.paragraph_format.space_after = Pt(4)
                run = cp.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)
                doc.add_paragraph()
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue
            
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 3. Markdown tables
        if line.startswith('|') and '|' in line[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
                
            if len(table_lines) >= 2:
                headers = [c.strip() for c in table_lines[0].split('|')[1:-1]]
                data_rows = []
                for row_line in table_lines[1:]:
                    if re.match(r'^\|[\s\-:|]+\|$', row_line):
                        continue
                    cols = [c.strip() for c in row_line.split('|')[1:-1]]
                    if cols:
                        data_rows.append(cols)
                        
                num_cols = len(headers)
                if num_cols > 0:
                    tbl = doc.add_table(rows=len(data_rows) + 1, cols=num_cols)
                    tbl.autofit = True
                    
                    # Style Header
                    hdr_cells = tbl.rows[0].cells
                    for col_idx, header_text in enumerate(headers):
                        if col_idx < len(hdr_cells):
                            hdr_cells[col_idx].text = header_text
                            set_cell_background(hdr_cells[col_idx], "2B6CB0")
                            p = hdr_cells[col_idx].paragraphs[0]
                            p.paragraph_format.space_before = Pt(4)
                            p.paragraph_format.space_after = Pt(4)
                            for r in p.runs:
                                r.font.bold = True
                                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                r.font.size = Pt(10)
                                
                    # Data Rows
                    for row_idx, row_data in enumerate(data_rows):
                        row_cells = tbl.rows[row_idx + 1].cells
                        bg_color = "F7FAFC" if row_idx % 2 == 1 else "FFFFFF"
                        for col_idx, cell_value in enumerate(row_data):
                            if col_idx < len(row_cells):
                                p = row_cells[col_idx].paragraphs[0]
                                p.text = "" # Clear
                                set_cell_background(row_cells[col_idx], bg_color)
                                p.paragraph_format.space_before = Pt(3)
                                p.paragraph_format.space_after = Pt(3)
                                render_paragraph_with_inline_math(p, cell_value)
                                for r in p.runs:
                                    if not r.font.size:
                                        r.font.size = Pt(9.5)
                                    
                    doc.add_paragraph()
            continue

        # Blank lines
        if not line.strip():
            i += 1
            continue

        # Headings
        if line.startswith('# '):
            h = doc.add_heading(level=1)
            run = h.add_run(line[2:].strip())
            run.font.name = 'Calibri'
            run.font.size = Pt(22)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(8)
            i += 1
            continue
        elif line.startswith('## '):
            h = doc.add_heading(level=2)
            run = h.add_run(line[3:].strip())
            run.font.name = 'Calibri'
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        elif line.startswith('### '):
            h = doc.add_heading(level=3)
            run = h.add_run(line[4:].strip())
            run.font.name = 'Calibri'
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(4)
            i += 1
            continue
        elif line.startswith('#### '):
            h = doc.add_heading(level=4)
            run = h.add_run(line[5:].strip())
            run.font.name = 'Calibri'
            run.font.size = Pt(11.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ['---', '***', '___']:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run('_______________________________________________________________________________')
            run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
            i += 1
            continue

        # Bullet list items
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            bullet_text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            render_paragraph_with_inline_math(p, bullet_text)
            i += 1
            continue

        # Numbered list items
        num_match = re.match(r'^(\d+)\.\s+(.*)$', line.strip())
        if num_match:
            item_text = num_match.group(2)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            render_paragraph_with_inline_math(p, item_text)
            i += 1
            continue

        # Blockquotes
        if line.strip().startswith('>'):
            quote_text = line.strip()[1:].strip()
            table = doc.add_table(rows=1, cols=1)
            table.autofit = False
            table.columns[0].width = Inches(6.5)
            cell = table.cell(0, 0)
            set_cell_background(cell, "EBF8FF")
            p = cell.paragraphs[0]
            p.paragraph_format.left_indent = Inches(0.1)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(quote_text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
            doc.add_paragraph()
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        render_paragraph_with_inline_math(p, line)
        i += 1

    doc.save(docx_path)
    print(f"Generated Beautiful Document with Math: {docx_path}")

if __name__ == "__main__":
    brain_dir = r"C:\Users\Manoj\.gemini\antigravity\brain\82a1cfbc-f612-4c89-a236-03801d02b37d"
    repo_dir = r"C:\genomicsecure"
    
    # 1. Non-Technical Project Report
    md1 = os.path.join(brain_dir, "project_report.md")
    markdown_to_docx_with_equations(md1, os.path.join(brain_dir, "MedLink_Project_Report.docx"), "Med-Link: Project Report")
    markdown_to_docx_with_equations(md1, os.path.join(repo_dir, "MedLink_Project_Report.docx"), "Med-Link: Project Report")
    
    # 2. Technical Architecture Reference Report
    md2 = os.path.join(brain_dir, "technical_architecture_report.md")
    markdown_to_docx_with_equations(md2, os.path.join(brain_dir, "MedLink_Technical_Architecture_Report.docx"), "Med-Link: Technical Architecture Reference")
    markdown_to_docx_with_equations(md2, os.path.join(repo_dir, "MedLink_Technical_Architecture_Report.docx"), "Med-Link: Technical Architecture Reference")
    
    # 3. In-Depth Backend Architecture & API Guide
    md3 = os.path.join(brain_dir, "backend_in_depth_report.md")
    markdown_to_docx_with_equations(md3, os.path.join(brain_dir, "MedLink_Backend_InDepth_Report.docx"), "Med-Link: In-Depth Backend Architecture & API Guide")
    markdown_to_docx_with_equations(md3, os.path.join(repo_dir, "MedLink_Backend_InDepth_Report.docx"), "Med-Link: In-Depth Backend Architecture & API Guide")
